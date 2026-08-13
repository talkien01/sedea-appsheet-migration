"""Genera el .docx de la ficha a partir del dict que produce ficha_engine.generar()."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x36, 0x60, 0x92)
RED = RGBColor(0xA6, 0x19, 0x2E)
REDFILL = "FDE9E9"

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_font(run, size=9, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def _header_row(table, headers, widths):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.width = Inches(widths[i])
        _shade(cell, "366092")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

def _data_row(table, values, widths, bold=False, fill=None, center=True):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.width = Inches(widths[i])
        if fill:
            _shade(cell, fill)
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(v))
        _set_font(run, size=9, bold=bold)
    return row

def _fmt(n):
    try:
        return "$" + format(round(float(n)), ",")
    except (TypeError, ValueError):
        return str(n) if n else ""

def escribir_ficha(data, out_path):
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(0.5)
        section.top_margin = section.bottom_margin = Inches(0.5)

    h = doc.add_paragraph()
    r = h.add_run("SECRETARÍA DE DESARROLLO AGROPECUARIO")
    _set_font(r, size=13, bold=True)

    h2 = doc.add_paragraph()
    r2 = h2.add_run(f"FICHA MUNICIPAL — {data['municipio']}")
    _set_font(r2, size=15, bold=True)

    h3 = doc.add_paragraph()
    r3 = h3.add_run("Generada automáticamente — fuente: analitica (Postgres) + Datos_referencia_manual_municipios.xlsx")
    _set_font(r3, size=9, bold=False)
    r3.italic = True

    doc.add_paragraph()

    # ---- Warnings ----
    if data["warnings"]:
        p = doc.add_paragraph()
        r = p.add_run(f"ADVERTENCIAS — {len(data['warnings'])} pendiente(s) antes de publicar")
        _set_font(r, size=11, bold=True, color=RED)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _header_row(table, ["#", "Advertencia — dato faltante y qué hacer"], [0.4, 6.6])
        for i, w in enumerate(data["warnings"]):
            _data_row(table, [i+1, w], [0.4, 6.6], fill=REDFILL, center=False)
        doc.add_paragraph()

    # ---- 1. Histórico ----
    p = doc.add_paragraph()
    r = p.add_run(f"1. Apoyos e inversión por programa, {'-'.join(data['anios_presentes']) if data['anios_presentes'] else 'sin datos'} (analitica.apoyo_municipio)")
    _set_font(r, size=11, bold=True, color=BLUE)

    if data["historico"]:
        # R3: la tabla muestra las 5 columnas de dinero, incluida Federal.
        widths = [0.5, 1.5, 0.7, 1.0, 1.0, 1.0, 1.2, 1.1]
        table = doc.add_table(rows=1, cols=8)
        table.autofit = False
        _header_row(table, ["Año","Programa","Apoyos","Federal","Apoyo Estatal","Apoyo Municipal","Aportación Productores","Total"], widths)
        for row_ in data["historico"]:
            _data_row(table, [row_["anio"], row_["programa_nombre"], row_["numero_apoyos"],
                               _fmt(row_.get("apoyo_federal")),
                               _fmt(row_["apoyo_estatal"]), _fmt(row_["apoyo_municipal"]),
                               _fmt(row_["aportacion_productor"]), _fmt(row_["total"])], widths)
        _data_row(table, ["", "TOTAL", data["total_apoyos_historico"], "", "", "", "", _fmt(data["total_inversion_historico"])], widths, bold=True)
    else:
        doc.add_paragraph("Sin datos históricos para este municipio.")
    doc.add_paragraph()

    # ---- 2. Avance 2026 ----
    p = doc.add_paragraph()
    r = p.add_run("2. Avance 2026 (analitica.v_oficial_municipio)")
    _set_font(r, size=11, bold=True, color=BLUE)
    if data["avance_2026"]:
        widths = [2.4, 1.1, 1.1, 1.4, 1.4]
        table = doc.add_table(rows=1, cols=5)
        table.autofit = False
        _header_row(table, ["Componente","Solicitudes","Apoyos","Estatal dictaminado","Total dictaminado"], widths)
        for row_ in data["avance_2026"]:
            _data_row(table, [row_["componente"], row_["solicitudes"], row_["apoyos"],
                               _fmt(row_["estatal_dictaminado"]), _fmt(row_["total_dictaminado"])], widths)
        _data_row(table, ["TOTAL", "", data["total_apoyos_2026"], "", _fmt(data["total_2026"])], widths, bold=True)
    else:
        doc.add_paragraph("Sin datos 2026 en v_oficial_municipio para este municipio.")
    doc.add_paragraph()

    # ---- 3-6. Bloques manuales ----
    bloques = [
        ("3. Territorio y superficie agrícola", data["territorio"],
         ["Municipio","Extensión territorial (Ha)","% del territorio estatal","Superficie agrícola total (Ha)","Riego (Ha)","Temporal (Ha)"]),
        ("4. Top de productos", data["productos"],
         ["Municipio","Rank","Producto","Superficie (Ha)","Volumen (Ton)","Valor (MDP)"]),
        ("5. Precipitación mensual", data["precipitacion"],
         ["Municipio","Año","ENE","FEB","MAR","ABR","MAY","JUN","JUL","AGO","SEP","OCT","NOV","DIC","Anual"]),
        ("6. Demografía municipal de beneficiarios", data["demografia"],
         ["Municipio","Grupo de edad","Hombres","Mujeres","Total","% del total municipal"]),
    ]
    for titulo, rows_, cols in bloques:
        p = doc.add_paragraph()
        r = p.add_run(titulo)
        _set_font(r, size=11, bold=True, color=BLUE)
        vacio = (not rows_) or all(
            all((v is None or v == "") for k, v in row_.items() if k not in ("Municipio",))
            for row_ in rows_
        )
        if vacio:
            pp = doc.add_paragraph()
            rr = pp.add_run("PENDIENTE — sin datos cargados en la plantilla manual todavía.")
            _set_font(rr, size=9, color=RED)
            rr.italic = True
        else:
            w = 7.0 / len(cols)
            widths = [w] * len(cols)
            table = doc.add_table(rows=1, cols=len(cols))
            table.autofit = False
            _header_row(table, cols, widths)
            for row_ in rows_:
                _data_row(table, [row_.get(c, "") if row_.get(c) is not None else "" for c in cols], widths)
        doc.add_paragraph()

    _secciones_extendidas(doc, data, desde=7)
    doc.save(out_path)


# ===========================================================================
# Extensión 2026: secciones 7-10 comunes y fichas regional / estatal.
# ===========================================================================
VACIO = "—"


def _pct(v):
    if v is None or v == "":
        return VACIO
    try:
        return f"{float(v):.1f}%"
    except (TypeError, ValueError):
        return VACIO


def _num(v):
    if v is None or v == "":
        return VACIO
    try:
        return f"{int(float(v)):,}"
    except (TypeError, ValueError):
        return str(v)


def _dinero(v):
    """Monto con formato; sin dato ⇒ «—», nunca $0 (R1/R4)."""
    if v is None or v == "":
        return VACIO
    try:
        return "$" + format(round(float(v)), ",")
    except (TypeError, ValueError):
        return VACIO


def _titulo(doc, texto_titulo):
    p = doc.add_paragraph()
    r = p.add_run(texto_titulo)
    _set_font(r, size=11, bold=True, color=BLUE)


def _leyenda(doc, texto_leyenda):
    p = doc.add_paragraph()
    r = p.add_run(texto_leyenda)
    _set_font(r, size=9, color=RED)
    r.italic = True


def _tabla(doc, encabezados, filas, anchos=None):
    anchos = anchos or [7.0 / len(encabezados)] * len(encabezados)
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.autofit = False
    _header_row(t, encabezados, anchos)
    for fila in filas:
        _data_row(t, fila, anchos)
    return t


def _secciones_extendidas(doc, data, desde=7):
    n = desde

    # --- Emergentes vs Productividad ---
    _titulo(doc, f"{n}. Distribución Emergentes / Productividad (analitica.vw_matriz_emer_prod)")
    filas = data.get("emer_prod") or []
    if not filas:
        _leyenda(doc, data.get("emer_prod_leyenda")
                 or "Sin datos de clasificación Emergentes / Productividad para este ámbito.")
    else:
        agrupado = {}
        for f in filas:
            k = (f.get("anio"), f.get("clasificacion"))
            a = agrupado.setdefault(k, {"apoyos": None, "total": None})
            for origen, destino in (("numero_apoyos", "apoyos"), ("total", "total")):
                v = f.get(origen)
                if v is not None:
                    a[destino] = (a[destino] or 0) + v
        renglones = [[anio, clas, _num(v["apoyos"]), _dinero(v["total"])]
                     for (anio, clas), v in sorted(agrupado.items(),
                                                   key=lambda kv: (kv[0][0] or 0, kv[0][1] or ""))]
        _tabla(doc, ["Año", "Clasificación (Emergentes / Productividad)", "Apoyos", "Inversión total"],
               renglones, [0.9, 3.3, 1.2, 1.6])
    doc.add_paragraph()
    n += 1

    # --- Aportaciones ---
    _titulo(doc, f"{n}. Distribución de aportaciones (Federal / Estatal / Municipal / Beneficiario)")
    ap = data.get("aportaciones") or {}
    if not ap:
        _leyenda(doc, data.get("aportaciones_leyenda")
                 or "Sin desglose de aportaciones para este ámbito.")
    else:
        renglones = [
            ["Federal", _dinero(ap.get("federal")), _pct(ap.get("pct_federal"))],
            ["Estatal", _dinero(ap.get("estatal")), _pct(ap.get("pct_estatal"))],
            ["Municipal", _dinero(ap.get("municipal")), _pct(ap.get("pct_municipal"))],
            ["Beneficiario (productor)", _dinero(ap.get("beneficiario")), _pct(ap.get("pct_beneficiario"))],
            ["TOTAL", _dinero(ap.get("total")), "100.0%" if ap.get("total") else VACIO],
        ]
        _tabla(doc, ["Aportación", "Monto", "% del total"], renglones, [2.6, 2.2, 2.2])
        _leyenda(doc, "Las celdas con «—» no son cero: son montos que la fuente no desglosa (R1/R3).")
    doc.add_paragraph()
    n += 1

    # --- Género y edad ---
    _titulo(doc, f"{n}. Género y edad de las personas beneficiarias (derivado de CURP)")
    filas = data.get("genero_edad") or []
    if not filas:
        _leyenda(doc, data.get("genero_edad_leyenda")
                 or "Sin CURP cargada para este ámbito: no se reporta género ni edad. "
                    "No se muestran ceros ni estimaciones (R1/R6).")
    else:
        renglones = [[f.get("rango_edad") or VACIO,
                      {"H": "Hombres", "M": "Mujeres"}.get(f.get("genero"), VACIO),
                      _num(f.get("personas"))] for f in filas]
        _tabla(doc, ["Rango de edad", "Género", "Personas (CURP distintas)"],
               renglones, [2.4, 2.2, 2.4])
        _leyenda(doc, "Personas con CURP válida y distinta; no es el número de apoyos (R2). "
                      "El sexo sale de la CURP, nunca del nombre (R6).")
    doc.add_paragraph()
    n += 1

    # --- Incidencias ---
    _titulo(doc, f"{n}. Incidencias abiertas de calidad de datos")
    filas = data.get("incidencias") or []
    if not filas:
        _leyenda(doc, data.get("incidencias_leyenda") or "Sin incidencias abiertas.")
    else:
        renglones = [[f.get("severidad"), f.get("tipo"), f.get("descripcion"),
                      f.get("accion_sugerida")] for f in filas[:25]]
        _tabla(doc, ["Severidad", "Tipo", "Descripción", "Acción sugerida"],
               renglones, [1.0, 1.5, 2.4, 2.1])
    doc.add_paragraph()


def _encabezado(doc, titulo_ficha, subtitulo):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(0.5)
        section.top_margin = section.bottom_margin = Inches(0.5)
    p = doc.add_paragraph()
    _set_font(p.add_run("SECRETARÍA DE DESARROLLO AGROPECUARIO"), size=13, bold=True)
    p2 = doc.add_paragraph()
    _set_font(p2.add_run(titulo_ficha), size=15, bold=True)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(subtitulo)
    _set_font(r3, size=9)
    r3.italic = True
    doc.add_paragraph()


def _bloque_advertencias(doc, warnings):
    if not warnings:
        return
    p = doc.add_paragraph()
    _set_font(p.add_run(f"ADVERTENCIAS — {len(warnings)} pendiente(s) antes de publicar"),
              size=11, bold=True, color=RED)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _header_row(table, ["#", "Advertencia — dato faltante y qué hacer"], [0.4, 6.6])
    for i, w in enumerate(warnings):
        _data_row(table, [i + 1, w], [0.4, 6.6], fill=REDFILL, center=False)
    doc.add_paragraph()


def escribir_ficha_region(data, out_path):
    """Ficha regional: mismas secciones que la municipal, agregadas por región."""
    doc = Document()
    _encabezado(doc, f"FICHA REGIONAL — REGIÓN {data['region']}",
                "Generada automáticamente desde analitica (Postgres). "
                "Los montos son totales (federal + estatal + municipal + productor).")
    _bloque_advertencias(doc, data.get("warnings"))

    _titulo(doc, "1. Municipios de la región y su acumulado histórico")
    renglones = []
    for m in data["resumen_municipios"]:
        renglones.append([m["municipio"], _num(m.get("numero_apoyos")), _dinero(m.get("federal")),
                          _dinero(m.get("estatal")), _dinero(m.get("municipal")),
                          _dinero(m.get("beneficiario")), _dinero(m.get("total"))])
    tot = data.get("total") or {}
    renglones.append(["TOTAL REGIÓN", _num(tot.get("numero_apoyos")), _dinero(tot.get("federal")),
                      _dinero(tot.get("estatal")), _dinero(tot.get("municipal")),
                      _dinero(tot.get("beneficiario")), _dinero(tot.get("total"))])
    _tabla(doc, ["Municipio", "Apoyos", "Federal", "Estatal", "Municipal",
                 "Beneficiario (productor)", "Total"], renglones,
           [1.5, 0.8, 1.0, 1.0, 1.0, 1.1, 1.1])
    doc.add_paragraph()

    _titulo(doc, "2. Inversión por año")
    serie = data.get("serie_anual") or []
    if not serie:
        _leyenda(doc, "Sin serie anual disponible.")
    else:
        _tabla(doc, ["Año", "Federal", "Estatal", "Municipal", "Beneficiario (productor)", "Total"],
               [[s.get("anio"), _dinero(s.get("federal")), _dinero(s.get("estatal")),
                 _dinero(s.get("municipal")), _dinero(s.get("beneficiario")), _dinero(s.get("total"))]
                for s in serie], [0.8, 1.2, 1.2, 1.2, 1.4, 1.2])
        _leyenda(doc, "Los años sin carga no aparecen: vacío no es cero (R4). "
                      "2027 está previsto y vacío por definición.")
    doc.add_paragraph()

    _secciones_extendidas(doc, data, desde=3)
    doc.save(out_path)


def escribir_ficha_estatal(data, out_path):
    doc = Document()
    _encabezado(doc, "FICHA ESTATAL — QUERÉTARO",
                "Generada automáticamente desde analitica (Postgres): resumen_estatal, "
                "v_oficial_componente y la matriz histórica.")
    _bloque_advertencias(doc, data.get("warnings"))

    _titulo(doc, "1. Inversión estatal por año (todas las aportaciones)")
    serie = data.get("serie_anual") or []
    if not serie:
        _leyenda(doc, "Sin serie anual disponible.")
    else:
        _tabla(doc, ["Año", "Apoyos", "Federal", "Estatal", "Municipal",
                     "Beneficiario (productor)", "Total"],
               [[s.get("anio"), _num(s.get("numero_apoyos")), _dinero(s.get("federal")),
                 _dinero(s.get("estatal")), _dinero(s.get("municipal")),
                 _dinero(s.get("beneficiario")), _dinero(s.get("total"))] for s in serie],
               [0.7, 0.8, 1.1, 1.1, 1.1, 1.2, 1.1])
    doc.add_paragraph()

    _titulo(doc, "2. Resumen estatal por programa (analitica.resumen_estatal)")
    filas = data.get("resumen_estatal") or []
    if not filas:
        _leyenda(doc, "Sin datos en resumen_estatal.")
    else:
        _tabla(doc, ["Año", "Programa", "Apoyos", "Federal", "Estatal", "Total"],
               [[f.get("anio"), f.get("programa_nombre"), _num(f.get("numero_apoyos")),
                 _dinero(f.get("apoyo_federal")), _dinero(f.get("apoyo_estatal")),
                 _dinero(f.get("total"))] for f in filas[:60]],
               [0.7, 2.4, 0.8, 1.0, 1.0, 1.1])
    doc.add_paragraph()

    _titulo(doc, "3. Componentes 2026 dictaminados (analitica.v_oficial_componente)")
    comps = data.get("componentes_2026") or []
    if not comps:
        _leyenda(doc, "Sin datos en v_oficial_componente.")
    else:
        encabezados = list(comps[0].keys())
        _tabla(doc, [h.replace("_", " ").title() for h in encabezados],
               [[c.get(h) if c.get(h) is not None else VACIO for h in encabezados] for c in comps])
    doc.add_paragraph()

    _secciones_extendidas(doc, data, desde=4)
    doc.save(out_path)
