"""Genera el .docx de la ficha a partir del dict que produce ficha_engine.generar()."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x36, 0x60, 0x92)
RED = RGBColor(0xA6, 0x19, 0x2E)
REDFILL = "FDE9E9"

def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_font(run, size=9, bold=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def _header_row(table, headers, widths):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.width = Inches(widths[i])
        _shade(cell, "366092")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

def _data_row(table, values, widths, bold=False, fill=None, center=True):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.width = Inches(widths[i])
        if fill:
            _shade(cell, fill)
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(v))
        _set_font(run, size=9, bold=bold)
    return row

def _fmt(n):
    try:
        return "$" + format(round(float(n)), ",")
    except (TypeError, ValueError):
        return str(n) if n else ""

def escribir_ficha(data, out_path):
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = section.right_margin = Inches(0.5)
        section.top_margin = section.bottom_margin = Inches(0.5)

    h = doc.add_paragraph()
    r = h.add_run("SECRETARÍA DE DESARROLLO AGROPECUARIO")
    _set_font(r, size=13, bold=True)

    h2 = doc.add_paragraph()
    r2 = h2.add_run(f"FICHA MUNICIPAL — {data['municipio']}")
    _set_font(r2, size=15, bold=True)

    h3 = doc.add_paragraph()
    r3 = h3.add_run("Generada automáticamente — fuente: analitica (Postgres) + Datos_referencia_manual_municipios.xlsx")
    _set_font(r3, size=9, bold=False)
    r3.italic = True

    doc.add_paragraph()

    # ---- Warnings ----
    if data["warnings"]:
        p = doc.add_paragraph()
        r = p.add_run(f"ADVERTENCIAS — {len(data['warnings'])} pendiente(s) antes de publicar")
        _set_font(r, size=11, bold=True, color=RED)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        _header_row(table, ["#", "Advertencia — dato faltante y qué hacer"], [0.4, 6.6])
        for i, w in enumerate(data["warnings"]):
            _data_row(table, [i+1, w], [0.4, 6.6], fill=REDFILL, center=False)
        doc.add_paragraph()

    # ---- 1. Histórico ----
    p = doc.add_paragraph()
    r = p.add_run(f"1. Apoyos e inversión por programa, {'-'.join(data['anios_presentes']) if data['anios_presentes'] else 'sin datos'} (analitica.apoyo_municipio)")
    _set_font(r, size=11, bold=True, color=BLUE)

    if data["historico"]:
        widths = [0.6, 1.8, 0.8, 1.1, 1.1, 1.3, 1.1]
        table = doc.add_table(rows=1, cols=7)
        table.autofit = False
        _header_row(table, ["Año","Programa","Apoyos","Apoyo Estatal","Apoyo Municipal","Aportación Productores","Total"], widths)
        for row_ in data["historico"]:
            _data_row(table, [row_["anio"], row_["programa_nombre"], row_["numero_apoyos"],
                               _fmt(row_["apoyo_estatal"]), _fmt(row_["apoyo_municipal"]),
                               _fmt(row_["aportacion_productor"]), _fmt(row_["total"])], widths)
        _data_row(table, ["", "TOTAL", data["total_apoyos_historico"], "", "", "", _fmt(data["total_inversion_historico"])], widths, bold=True)
    else:
        doc.add_paragraph("Sin datos históricos para este municipio.")
    doc.add_paragraph()

    # ---- 2. Avance 2026 ----
    p = doc.add_paragraph()
    r = p.add_run("2. Avance 2026 (analitica.v_oficial_municipio)")
    _set_font(r, size=11, bold=True, color=BLUE)
    if data["avance_2026"]:
        widths = [2.4, 1.1, 1.1, 1.4, 1.4]
        table = doc.add_table(rows=1, cols=5)
        table.autofit = False
        _header_row(table, ["Componente","Solicitudes","Apoyos","Estatal dictaminado","Total dictaminado"], widths)
        for row_ in data["avance_2026"]:
            _data_row(table, [row_["componente"], row_["solicitudes"], row_["apoyos"],
                               _fmt(row_["estatal_dictaminado"]), _fmt(row_["total_dictaminado"])], widths)
        _data_row(table, ["TOTAL", "", data["total_apoyos_2026"], "", _fmt(data["total_2026"])], widths, bold=True)
    else:
        doc.add_paragraph("Sin datos 2026 en v_oficial_municipio para este municipio.")
    doc.add_paragraph()

    # ---- 3-6. Bloques manuales ----
    bloques = [
        ("3. Territorio y superficie agrícola", data["territorio"],
         ["Municipio","Extensión territorial (Ha)","% del territorio estatal","Superficie agrícola total (Ha)","Riego (Ha)","Temporal (Ha)"]),
        ("4. Top de productos", data["productos"],
         ["Municipio","Rank","Producto","Superficie (Ha)","Volumen (Ton)","Valor (MDP)"]),
        ("5. Precipitación mensual", data["precipitacion"],
         ["Municipio","Año","ENE","FEB","MAR","ABR","MAY","JUN","JUL","AGO","SEP","OCT","NOV","DIC","Anual"]),
        ("6. Demografía municipal de beneficiarios", data["demografia"],
         ["Municipio","Grupo de edad","Hombres","Mujeres","Total","% del total municipal"]),
    ]
    for titulo, rows_, cols in bloques:
        p = doc.add_paragraph()
        r = p.add_run(titulo)
        _set_font(r, size=11, bold=True, color=BLUE)
        vacio = (not rows_) or all(
            all((v is None or v == "") for k, v in row_.items() if k not in ("Municipio",))
            for row_ in rows_
        )
        if vacio:
            pp = doc.add_paragraph()
            rr = pp.add_run("PENDIENTE — sin datos cargados en la plantilla manual todavía.")
            _set_font(rr, size=9, color=RED)
            rr.italic = True
        else:
            w = 7.0 / len(cols)
            widths = [w] * len(cols)
            table = doc.add_table(rows=1, cols=len(cols))
            table.autofit = False
            _header_row(table, cols, widths)
            for row_ in rows_:
                _data_row(table, [row_.get(c, "") if row_.get(c) is not None else "" for c in cols], widths)
        doc.add_paragraph()

    doc.save(out_path)
